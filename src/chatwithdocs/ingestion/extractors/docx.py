from __future__ import annotations

import logging
import uuid
from pathlib import Path
from typing import Any

from chatwithdocs.ingestion.base import Chunk, ExtractionResult, Extractor, FileType

logger = logging.getLogger(__name__)


class DOCXExtractor(Extractor):
    """Extract text and tables from DOCX files using python-docx.

    This extractor handles:
    - Text extraction from paragraphs
    - Table extraction (converted to markdown format)
    - Document structure tracking (headings as section markers)
    - Appropriate text chunking
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract content from DOCX file.

        Args:
            file_path: Path to DOCX file.

        Returns:
            ExtractionResult with chunks and metadata.
        """
        try:
            import docx
        except ImportError:
            error_msg = "python-docx not installed. Install with: uv pip install python-docx"
            logger.error(error_msg)
            return ExtractionResult(errors=[error_msg])

        chunks: list[Chunk] = []
        errors: list[str] = []
        metadata: dict[str, Any] = {}

        try:
            doc = docx.Document(file_path)

            # Extract document metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
            }

            # Track current section header
            current_section = None

            # Process paragraphs with section tracking
            text_buffer = []
            buffer_size = 0

            for para in doc.paragraphs:
                try:
                    # Check if this is a heading
                    if para.style.name.startswith("Heading"):
                        # Flush current buffer as a chunk if we have content
                        if text_buffer:
                            chunk_text = " ".join(text_buffer)
                            page_chunks = self._chunk_text(chunk_text, file_path, current_section)
                            chunks.extend(page_chunks)
                            text_buffer = []
                            buffer_size = 0

                        # Update section header
                        current_section = para.text.strip()
                        continue

                    text = para.text.strip()
                    if text:
                        text_buffer.append(text)
                        buffer_size += len(text)

                        # Flush buffer when it exceeds chunk size
                        if buffer_size >= self.chunk_size:
                            chunk_text = " ".join(text_buffer)
                            page_chunks = self._chunk_text(chunk_text, file_path, current_section)
                            chunks.extend(page_chunks)

                            # Keep overlap for next chunk
                            overlap_text = self._get_overlap_text(text_buffer)
                            text_buffer = [overlap_text] if overlap_text else []
                            buffer_size = len(overlap_text)

                except Exception as e:
                    error_msg = f"Error processing paragraph: {e}"
                    logger.warning(error_msg)
                    errors.append(error_msg)

            # Don't forget remaining text in buffer
            if text_buffer:
                chunk_text = " ".join(text_buffer)
                page_chunks = self._chunk_text(chunk_text, file_path, current_section)
                chunks.extend(page_chunks)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_chunk = self._table_to_chunk(table, file_path, table_idx, current_section)
                    chunks.append(table_chunk)
                except Exception as e:
                    error_msg = f"Error extracting table {table_idx}: {e}"
                    logger.warning(error_msg)
                    errors.append(error_msg)

        except Exception as e:
            error_msg = f"Error opening DOCX {file_path}: {e}"
            logger.error(error_msg)
            errors.append(error_msg)

        return ExtractionResult(chunks=chunks, metadata=metadata, errors=errors)

    def _chunk_text(self, text: str, source_file: Path, section_header: str | None) -> list[Chunk]:
        """Split text into overlapping chunks.

        Args:
            text: Text content to chunk.
            source_file: Source file path.
            section_header: Current section header.

        Returns:
            List of chunks.
        """
        chunks: list[Chunk] = []

        # Sentence-based chunking
        sentences = text.replace("\n", " ").split(". ")
        current_chunk = []
        current_size = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if not sentence.endswith("."):
                sentence += "."

            sentence_size = len(sentence)

            if current_size + sentence_size > self.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append(
                    Chunk(
                        id=str(uuid.uuid4()),
                        content=chunk_text,
                        source_file=source_file,
                        file_type=FileType.DOCX,
                        section_header=section_header,
                        chunk_type="text",
                    )
                )

                # Keep overlap sentences
                overlap_size = 0
                overlap_sentences = []
                for s in reversed(current_chunk):
                    if overlap_size + len(s) > self.chunk_overlap:
                        break
                    overlap_sentences.insert(0, s)
                    overlap_size += len(s)

                current_chunk = overlap_sentences
                current_size = overlap_size

            current_chunk.append(sentence)
            current_size += sentence_size

        # Final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(
                Chunk(
                    id=str(uuid.uuid4()),
                    content=chunk_text,
                    source_file=source_file,
                    file_type=FileType.DOCX,
                    section_header=section_header,
                    chunk_type="text",
                )
            )

        return chunks

    def _get_overlap_text(self, text_buffer: list[str]) -> str:
        """Get overlap text from buffer for next chunk."""
        overlap_size = 0
        overlap_sentences = []

        for s in reversed(text_buffer):
            if overlap_size + len(s) > self.chunk_overlap:
                break
            overlap_sentences.insert(0, s)
            overlap_size += len(s)

        return " ".join(overlap_sentences)

    def _table_to_chunk(
        self,
        table: Any,
        source_file: Path,
        table_idx: int,
        section_header: str | None,
    ) -> Chunk:
        """Convert table to markdown format chunk.

        Args:
            table: docx Table object.
            source_file: Source file path.
            table_idx: Table index.
            section_header: Current section header.

        Returns:
            Chunk with table in markdown format.
        """
        lines = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

            if i == 0:
                separators = ["---"] * len(cells)
                lines.append("| " + " | ".join(separators) + " |")

        table_text = "\n".join(lines)

        return Chunk(
            id=str(uuid.uuid4()),
            content=f"Table:\n{table_text}",
            source_file=source_file,
            file_type=FileType.DOCX,
            section_header=section_header,
            chunk_type="table",
            metadata={"table_index": table_idx},
        )

    def supports(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return file_path.suffix.lower() == ".docx"

    @property
    def file_type(self) -> FileType:
        """Return DOCX file type."""
        return FileType.DOCX
